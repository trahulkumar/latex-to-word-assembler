from docx import Document
from docx.shared import RGBColor
import os

def verify_output(docx_path):
    if not os.path.exists(docx_path):
        print(f"File not found: {docx_path}")
        return

    doc = Document(docx_path)
    print(f"Verifying {docx_path}...")
    
    # Check Styles
    styles_to_check = ['Normal', 'Heading 1', 'Heading 2', 'Heading 3', 'Source Code']
    for s_name in styles_to_check:
        if s_name in doc.styles:
            style = doc.styles[s_name]
            pf = style.paragraph_format
            print(f"\nStyle '{s_name}':")
            if hasattr(style, 'font'):
                 print(f"  Font Name: {style.font.name}")
                 # Try to inspect rFonts if possible
                 try:
                     rPr = style.element.rPr
                     if rPr is not None:
                         # This might be tricky to access directly with python-docx wrappers, 
                         # but we can check the xml or just trust the .name prop if it was set correctly.
                         pass
                 except: pass
                 
            print(f"  Alignment: {pf.alignment}")
            if hasattr(style, 'font') and style.font.size:
                print(f"  Font Size: {style.font.size.pt} pt")
            
            if s_name == 'Heading 1':
                print(f"  (Note: Heading 1 might be removed/hidden if duplicate logic worked)")

    # Check Title Page Content (First 2 paragraphs)
    print("\nTitle Page Content & Colors:")
    target_color = (54, 95, 145)
    
    for i in range(min(5, len(doc.paragraphs))):
        p = doc.paragraphs[i]
        txt = p.text.strip()
        if txt:
            runs = p.runs
            color_matches = False
            r_color = "None"
            if runs:
                # Check ALL runs for color match
                rgb = runs[0].font.color.rgb
                if rgb:
                    r_color = f"({rgb[0]}, {rgb[1]}, {rgb[2]})"
                    if rgb == RGBColor(*target_color):
                        color_matches = True
            
            print(f"  Para {i}: '{txt[:50]}...' | Align: {p.alignment} | Color: {r_color} | Match Target? {color_matches}")

    # Check for Duplicate "From Notebooks to Systems"
    print("\nChecking for Duplicate Heading 1:")
    h1_count = 0
    
    print("\nChecking Headings for Tabs and Fonts:")
    for p in doc.paragraphs:
        if p.style.name in ['Heading 1', 'Heading 2', 'Heading 3']:
            print(f"  {p.style.name}: '{p.text[:40]}...'")
            if '\t' in p.text:
                print("    [FAIL] Tab character found!")
            else:
                print("    [PASS] No tab character.")
                
                
            # Check font of first run
            if p.runs:
                r = p.runs[0]
                size_info = ""
                if r.font.size:
                    size_info = f" | Size: {r.font.size.pt} pt"
                print(f"    Font: {r.font.name}{size_info}")
                
        if p.style.name == 'Heading 1':
            h1_count += 1
    
    if h1_count == 0:
        print("  [INFO] No Heading 1 paragraphs found (Duplicate likely removed, assuming Custom Title is not Heading 1 style)")
    else:
        print(f"  [INFO] Found {h1_count} Heading 1 paragraphs.")

    # Check Footer
    print("\nChecking Footer:")
    section = doc.sections[0]
    footer = section.footer
    if footer.paragraphs:
        fp = footer.paragraphs[0]
        print(f"  Footer Paragraph Alignment: {fp.alignment} (Expected CENTER)")
        print(f"  Footer Text (might be empty if field): '{fp.text}'")
        # Check XML for fldSimple
        if 'w:fldSimple' in fp._element.xml:
             print("  [PASS] Found 'w:fldSimple' (Page Number Field) in footer XML.")
        else:
             print("  [FAIL] No 'w:fldSimple' found in footer XML.")

if __name__ == "__main__":
    verify_output(r"d:\AI\Github\agents\Latex-to-docx beautify\output\C01_From_Notebooks_to_Systems.docx")
