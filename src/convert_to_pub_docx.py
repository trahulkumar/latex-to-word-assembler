import json
import os
import pypandoc
import sys
import re
import argparse
import tempfile

def convert_book(metadata_path, output_dir, target_chapter=None):
    if not os.path.exists(metadata_path):
        print(f"Error: Metadata file not found at {metadata_path}")
        return

    with open(metadata_path, 'r', encoding='utf-8') as f:
        book_data = json.load(f)

    book_title = book_data.get("book_title", "Book")
    chapters = book_data.get("chapters", [])
    
    if not chapters:
        print("No chapters found in metadata.")
        return

    # Filter chapters if target_chapter is specified
    if target_chapter is not None:
        chapters = [c for c in chapters if c['number'] == target_chapter]
        if not chapters:
            print(f"Error: Chapter {target_chapter} not found in metadata.")
            return

    # Base directory for latex files
    base_latex_dir = "input/latex_files"

    for chapter in chapters:
        chapter_num = chapter['number']
        chapter_title = chapter['title']
        print(f"Processing Chapter {chapter_num}: {chapter_title}")
        
        # Expected chapter directory
        chapter_dir = os.path.join(base_latex_dir, f"Chapter_{chapter_num}")
        
        if not os.path.exists(chapter_dir):
            print(f"Error: Chapter directory not found: {chapter_dir}")
            sys.exit(1)

        # List all files in the chapter directory once
        try:
            dir_files = os.listdir(chapter_dir)
        except OSError as e:
            print(f"Error accessing directory {chapter_dir}: {e}")
            sys.exit(1)

        chapter_files = []
        for section in chapter.get('sections', []):
            section_num = section.get('number') # e.g. "1.1"
            
            found_file = None
            pattern = re.compile(rf"section[\s._]*{re.escape(str(section_num))}(\D|$)", re.IGNORECASE)
            
            for fname in dir_files:
                if pattern.search(fname):
                    found_file = os.path.join(chapter_dir, fname)
                    break 
            
            if not found_file:
                print(f"ERROR: Missing file for Section {section_num} in {chapter_dir}")
                print(f"       Expected file containing 'Section {section_num}'")
                sys.exit(1)
                
            print(f"  Found: {os.path.basename(found_file)}")
            chapter_files.append(found_file)
            
        if not chapter_files:
            print(f"  No valid files found for Chapter {chapter_num}")
            continue

        # Pre-process files
        temp_files = []
        cleaned_chapter_files = []
        
        citation_pattern = re.compile(r'\\(cite|citep|citet|ref)\{[^}]+\}|\[cite:[^\]]+\]|\[cite_start\]')
        graphics_pattern = re.compile(r'\\includegraphics(?:\[(.*?)\])?\{(.*?)\}')
        
        # Store references for consolidation
        references = {}


        # Publisher Style Patterns
        regex_eg = re.compile(r'\be\.g\.', re.IGNORECASE)
        regex_vs = re.compile(r'\bvs\.', re.IGNORECASE)
        regex_title_colon = re.compile(r'\\(section|subsection|subsubsection|paragraph)\{([^}]+):\s*\}')
        regex_caption_period = re.compile(r'(\\caption\{((?:[^{}]|\{[^{}]*\})*))\.\s*\}')
        regex_fig_ref_explicit = re.compile(r'\b(Figure|Table)\s+(\d+\.\d+)')
        regex_fig_ref_latex = re.compile(r'\b(Figure|Table)(~|\s+)(\\ref\{[^}]+\})')

        available_images = {}
        # heuristic: match on "figure_d_d"
        prefix_map = {}
        prefix_pattern = re.compile(r'^(figure_\d+_\d+)')

        for f in dir_files:
            name, ext = os.path.splitext(f)
            lower_name = name.lower()
            available_images[lower_name] = f
            
            pmatch = prefix_pattern.match(lower_name)
            if pmatch:
                prefix = pmatch.group(1)
                prefix_map[prefix] = f

        def find_target_image(ref_path):
            ref_basename = os.path.basename(ref_path)
            ref_name_no_ext, _ = os.path.splitext(ref_basename)
            lower_ref_name = ref_name_no_ext.lower()
            
            target = available_images.get(lower_ref_name)
            if not target:
                pmatch = prefix_pattern.match(lower_ref_name)
                if pmatch and pmatch.group(1) in prefix_map:
                    target = prefix_map[pmatch.group(1)]
            return target

        def latex_escape(text):
            if not text: return ""
            # Escaping backslashes first, then other problematic characters
            res = text.replace('\\', r'\textbackslash ')
            res = res.replace('{', r'\{').replace('}', r'\}')
            res = res.replace('&', r'\&').replace('$', r'\$').replace('_', r'\_')
            return res

        figure_block_pattern = re.compile(
            r'(\\begin\{figure\}(?:.|\n)*?\\end\{figure\})(\s*(?:%.*(?:\n|$))*)', 
            re.IGNORECASE
        )

        def process_figure_block(match):
            full_block = match.group(1)
            
            # Extract Placeholder Title
            placeholder_text = "Unknown Placeholder"
            placeholder_match = re.search(r'\\textbf\{Figure Placeholder:\s*(.*?)\}', full_block, re.IGNORECASE)
            if placeholder_match:
                placeholder_text = placeholder_match.group(1).strip()
            
            # Extract Caption
            caption = ""
            caption_start = full_block.find(r'\caption{')
            if caption_start != -1:
                # Manual brace counting to handle nested braces like \texttt{...}
                content_start = caption_start + len(r'\caption{')
                brace_count = 1
                current_pos = content_start
                while brace_count > 0 and current_pos < len(full_block):
                    if full_block[current_pos] == '{':
                        brace_count += 1
                    elif full_block[current_pos] == '}':
                        brace_count -= 1
                    current_pos += 1
                
                if brace_count == 0:
                    caption = full_block[content_start:current_pos-1].strip()
                    # Basic cleanup
                    caption = caption.replace('\n', ' ').replace('  ', ' ')

            
            # Extract Label
            label = ""
            label_match = re.search(r'\\label\{(.*?)\}', full_block)
            if label_match:
                label = label_match.group(1).strip()

            # Extract Prompt Comments
            # Find all lines starting with % inside the block
            prompts = []
            for line in full_block.split('\n'):
                 if line.strip().startswith('%'):
                     prompts.append(line.strip())
            prompt_text = "\n".join(prompts)


            g_match = graphics_pattern.search(full_block)
            
            # Logic: If image exists, show image AND details? 
            # User request: "Bring the entire figure section ... into the final latex and word file... show it in red"
            # This implies they want to see the metadata (prompt, placeholder) even if the image exists? 
            # Or is this specifically for the placeholder case? 
            # The user's snippet shows a placeholder block. 
            # Let's assume this is for ANY figure block that matches our pattern, likely mostly placeholders.
            # But if a real image is there, we probably still want the image + details if requested?
            # Actually, standard behavior is Image + Caption. 
            # The User's request specifically cites the *placeholder* block example.
            # So I will prioritize the placeholder text if found. 
            
            # Construct the Red Block text.
            # We prefix with [FIGURE DETAIL] so post-processing can color it red.
            # We use Markdown bold ** for keys.
            
            # Attempt to extract figure number from filename or label
            figure_num_str = "X.Y"
            if g_match:
                 ref_path = g_match.group(2)
                 # Try fig_4_1_... -> 4.1
                 fn_match = re.search(r'fig_(\d+)_(\d+)', os.path.basename(ref_path), re.IGNORECASE)
                 if fn_match:
                     figure_num_str = f"{fn_match.group(1)}.{fn_match.group(2)}"
            
            # If no filename match, try label: fig:4.1 or fig:chapter_04_01
            if figure_num_str == "X.Y" and label:
                 l_match = re.search(r'(\d+)[._](\d+)', label)
                 if l_match:
                     figure_num_str = f"{l_match.group(1)}.{l_match.group(2)}"

            # Add the special caption line requested:
            # "Figure 4.1 - Silent failure versus defensive gate..."
            # This should follow the [FIGURE DETAIL] block? Or be part of it?
            # User said: "Right after ...[MISSING IMAGE]... i want to add a caption... [FIGURE DETAIL] Caption: ..."
            # And then: "This should show up as 'Figure 4.1 - ...'"
            # So we create a specific line for it.
            
            details_block = (
                f"\n\n[FIGURE DETAIL] **Figure Placeholder:** {placeholder_text}\n"
                f"[FIGURE DETAIL] **Ref Label:** {label}\n"
                f"[FIGURE DETAIL] **Prompt Information:**\n"
            )
            
            # Process prompt lines to be distinctive
            for p in prompts:
                # Escape the whole line content using our helper
                # p starts with % usually, e.g. "% GitHub Script Prompt..."
                clean_p = p.lstrip('%').strip()
                escaped_p = latex_escape(clean_p)
                details_block += f"[FIGURE DETAIL] \\% {escaped_p}\n"

            # Add the Caption with specific formatting marker
            if caption:
                # Use our manual brace counting result, but maybe clean it?
                # Captions are already formatted, so we don't escape the whole thing, 
                # but we do want the and [FIGURE CAPTION] prefix.
                details_block += f"\n[FIGURE CAPTION] Figure {figure_num_str} - {caption}\n\n"
            else:
                details_block += "\n"

            # If an image exists, we might want to show it too? 
            # If it's a true placeholder block (as in the example), it likely has a PLACEHOLDER image or fbox.
            # If we replace the whole block with text, we lose the fbox, which is fine as the text covers it.
            # If there is a real \includegraphics, we should probably keep it and append details?
            # But the request says "Bring the entire figure section... into the final latex... show it in red".
            # Loops like they want the Source/Metadata visible.
            
            # Process trailing comments (group 2)
            trailing_block = ""
            if len(match.groups()) > 1 and match.group(2):
                extra_comments = match.group(2).strip().split('\n')
                for c in extra_comments:
                    if c.strip().startswith('%'):
                        clean_c = c.strip().lstrip('%').strip()
                        secure_c = latex_escape(clean_c)
                        trailing_block += f"\n\n[GITHUB SCRIPT PROMPT] \\% {secure_c}"

            if g_match:
                 options = g_match.group(1)
                 ref_path = g_match.group(2)
                 target_file = find_target_image(ref_path)
                 if target_file:
                     # It's a real image. Return the FULL original block with the updated path.
                     # AND append the trailing red comments if any.
                     new_tag = f'\\includegraphics[{options}]{{{target_file}}}' if options else f'\\includegraphics{{{target_file}}}'
                     s, e = g_match.span()
                     new_block = full_block[:s] + new_tag + full_block[e:]
                     return new_block + trailing_block
                 else:
                     # Missing image - include details + trailing comments
                     return f"\n\n**[MISSING IMAGE: {ref_path}]**\n{details_block}{trailing_block}"

            
            # If no graphics match (just fbox/text placeholder), return details + trailing
            return details_block + trailing_block


        def resolve_inline_image(match):
            options = match.group(1)
            ref_path = match.group(2)
            target = find_target_image(ref_path)
            if target:
                 return f'\\includegraphics[{options}]{{{target}}}' if options else f'\\includegraphics{{{target}}}'
            else:
                 return match.group(0)

        for file_path in chapter_files:
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                
                cleaned_content = citation_pattern.sub('', content)
                cleaned_content = figure_block_pattern.sub(process_figure_block, cleaned_content)
                cleaned_content = graphics_pattern.sub(resolve_inline_image, cleaned_content)

                # NEW: Handle Listing Captions and Script Prompts
                # 1. Regex replace Script Prompts (Multi-line)
                # Capture "% GitHub Script Prompt" and any subsequent lines starting with "%"
                script_prompt_pattern = re.compile(r'(^\s*% GitHub Script Prompt.*(?:\n\s*%.*)*)', re.MULTILINE)
                
                def format_script_prompt(m):
                    block = m.group(1)
                    lines = block.split('\n')
                    formatted_lines = []
                    for line in lines:
                        # Strip leading % and whitespace
                        clean_content = line.strip().lstrip('%').strip()
                        # ESCAPE for safe LaTeX inclusion
                        secure_content = latex_escape(clean_content)
                        # Output as a visible paragraph with marker
                        # We add \% so it looks like a comment in the final doc, but is visible text
                        formatted_lines.append(f"\n\n[GITHUB SCRIPT PROMPT] \\% {secure_content}")
                    return "".join(formatted_lines)

                cleaned_content = script_prompt_pattern.sub(format_script_prompt, cleaned_content)
                
                # 2. Regex replace Listing Captions (inside block -> move out)
                listing_pattern = re.compile(r'(\\begin\{lstlisting\}.*?)(\\end\{lstlisting\})', re.DOTALL)
                
                def place_listing_caption(m):
                    content = m.group(1)
                    end_tag = m.group(2)
                    # Find lines matching # Full Listing in the block content
                    lines = content.split('\n')
                    new_lines = []
                    extracted_caption = None
                    
                    for line in lines:
                        if line.strip().startswith('# Full Listing'):
                            extracted_caption = line.strip()
                            continue # Remove from block
                        new_lines.append(line)
                    
                    new_block = "\n".join(new_lines)
                    # Use original tags if they were there (content starts with \begin{lstlisting})
                    # But actually 'content' group 1 is (\begin{lstlisting}.*?)
                    replacement = f"{new_block}{end_tag}"
                    if extracted_caption:
                        # Strip leading # and whitespace
                        clean_caption = extracted_caption.lstrip('#').strip()
                        # Escape problematic characters in caption
                        safe_caption = latex_escape(clean_caption)
                        replacement += f"\n\n[LISTING CAPTION] {safe_caption}\n"
                    return replacement

                cleaned_content = listing_pattern.sub(place_listing_caption, cleaned_content)

                # Reference Extraction
                
                # 1. First extract bibtex blocks
                bib_block_pattern = re.compile(r'(\\begin\{thebibliography\}\{.*?\}(.*?)\\end\{thebibliography\})', re.DOTALL)
                bib_item_pattern = re.compile(r'\\bibitem\{([^}]+)\}(.*?)(?=\\bibitem|\Z)', re.DOTALL)

                def process_bib_block(match):
                    block_content = match.group(2)
                    items = bib_item_pattern.findall(block_content)
                    for key, text in items:
                        # Clean up text (remove newlines, extra spaces)
                        clean_text = " ".join(text.split()).strip()
                        if clean_text not in references.values():
                            references[key] = clean_text
                    return "" # Remove the block from the file

                cleaned_content = bib_block_pattern.sub(process_bib_block, cleaned_content)

                # 2. Extract plain-text references following a References header
                # We look for the References header and take everything after it.
                # Headers might be preceded by \vspace, \newpage, etc.
                ref_header_pattern = re.compile(
                    r'(?:\\vspace\{[^}]+\}\s*)?(?:\\noindent\s*)?(?:\\textbf\{References\}|\\(?:section|subsection|subsubsection)\*?\{References\})\s*(.*)', 
                    re.IGNORECASE | re.DOTALL
                )
                
                ref_match = ref_header_pattern.search(cleaned_content)
                if ref_match:
                    ref_text_block = ref_match.group(1)
                    # The content before the match is kept
                    cleaned_content = cleaned_content[:ref_match.start()].strip()
                    
                    # Process lines in the reference block
                    ref_lines = ref_text_block.split('\n')
                    for line in ref_lines:
                        line = line.strip()
                        if line:
                            # Deduplicate by content
                            if line not in references.values():
                                ref_key = f"text_ref_{len(references)}"
                                references[ref_key] = line
                else:
                    # Fallback if no explicit header
                    lines = cleaned_content.split('\n')
                    potential_refs = []
                    for i, line in enumerate(lines):
                        # Relaxed heuristic: Author list or Corp Author followed by (Year)
                        if re.match(r'^((?:[A-Z][a-zA-Z\s,&\.]+|[A-Z][a-z]+,?\s+[A-Z]\.?)\s*\(\d{4}\))', line.strip()):
                            potential_refs.append(i)
                    if potential_refs:
                        first_ref_idx = potential_refs[0]
                        if first_ref_idx > len(lines) * 0.5:
                            for idx in range(first_ref_idx, len(lines)):
                                ref_text = lines[idx].strip()
                                if ref_text:
                                    if ref_text not in references.values():
                                        ref_key = f"text_ref_{len(references)}"
                                        references[ref_key] = ref_text
                            cleaned_content = "\n".join(lines[:first_ref_idx])


                # Publisher Style Replacements
                cleaned_content = regex_eg.sub("for example", cleaned_content)
                cleaned_content = regex_vs.sub("versus", cleaned_content)
                cleaned_content = regex_title_colon.sub(r'\\\1{\2}', cleaned_content)
                cleaned_content = regex_caption_period.sub(r'\1}', cleaned_content)
                
                # Italicize Figure/Table references
                # Explicit: Figure 1.1 -> \textit{Figure 1.1}
                cleaned_content = regex_fig_ref_explicit.sub(lambda m: f"\\textit{{{m.group(1)} {m.group(2)}}}", cleaned_content)
                # Latex Ref: Figure~\ref{...} -> \textit{Figure~\ref{...}}
                cleaned_content = regex_fig_ref_latex.sub(lambda m: f"\\textit{{{m.group(1)}{m.group(2)}{m.group(3)}}}", cleaned_content)

                cleaned_content = cleaned_content.replace("``", '"').replace("''", '"')
                cleaned_content = cleaned_content.replace("—", "-")
                cleaned_content = cleaned_content.replace("**", "")

                # Fix stripped block math delimiters
                cleaned_content = re.sub(r'(?m)^\s*\[\s*$', r'\\[', cleaned_content)
                cleaned_content = re.sub(r'(?m)^\s*\]\s*$', r'\\]', cleaned_content)

                fd, temp_path = tempfile.mkstemp(suffix='.tex', text=True)
                with os.fdopen(fd, 'w', encoding='utf-8') as tf:
                    tf.write(cleaned_content)
                
                cleaned_chapter_files.append(temp_path)
                temp_files.append(temp_path)
                
                print(f"    Processed {os.path.basename(file_path)}: {len(cleaned_content)} chars")
                
            except Exception as e:
                print(f"  Error processing file {file_path}: {e}")
                for t in temp_files: 
                    try: os.remove(t) 
                    except: pass
                continue

        # Custom Title Page for Publisher Style
        # Right aligned, specific text structure to easily style in post-processing
        title_content = (
            f"\\begin{{flushright}}\n"
            f"CHAPTER {chapter_num}\n"
            f"\\par\n" 
            f"\\vspace{{0.5cm}}\n"
            f"{chapter_title}\n"
            f"\\end{{flushright}}\n"
            f"\\thispagestyle{{empty}}\n" 
            f"\\newpage\n"
            f"\\tableofcontents\n"
            f"\\newpage\n"
            f"\\chapter{{{chapter_title}}}\n" 
        )

        fd_title, title_path = tempfile.mkstemp(suffix='.tex', text=True)
        with os.fdopen(fd_title, 'w', encoding='utf-8') as tf:
            tf.write(title_content)
        temp_files.append(title_path)
        
        # Consolidated Bibliography
        if references:
            print(f"  Consolidating {len(references)} unique references...")
            bib_content = "\n\\newpage\n\\section*{References}\n\\begin{description}\n"
            
            # Sort references by key or appearance? Let's sort simply by key for stability, or keep order?
            # Creating a simple list. O'Reilly style usually list them.
            # Using description list to handle the lack of automatic numbering if we want, 
            # OR standard bibliography.
            # Let's use a standard itemize or description since we stripped the \bibitem wrapper.
            # Actually, standard latex bibliography is easier if we want that look.
            
            bib_content = "\n\\newpage\n\\section*{References}\n\\begin{itemize}\n"
            
            for key, text in references.items():
                # O'Reilly style: usually just the text. 
                # If text starts with Author, formatted.
                bib_content += f"\\item {text}\n"
            
            bib_content += "\\end{itemize}\n"
            
            fd_bib, bib_path = tempfile.mkstemp(suffix='.tex', text=True)
            with os.fdopen(fd_bib, 'w', encoding='utf-8') as tf:
                tf.write(bib_content)
            temp_files.append(bib_path)
            # Add to final files
            cleaned_chapter_files.append(bib_path)

        final_files = [title_path] + cleaned_chapter_files
        
        sanitized_title = "".join(c for c in chapter_title if c.isalnum() or c in (' ', '_', '-')).strip()
        sanitized_title = sanitized_title.replace(" ", "_")
        
        output_filename = f"C{chapter_num:02d}_{sanitized_title}.docx"
        output_path = os.path.join(output_dir, output_filename)
        
        abs_chapter_dir = os.path.abspath(chapter_dir)
        resource_path = f"{abs_chapter_dir};{os.path.join(abs_chapter_dir, 'images')}"
        
        extra_args = [
            f'--resource-path={resource_path}',
            '--top-level-division=chapter'
        ]
        
        print(f"  Combining {len(final_files)} files into {output_path}...")
        
        debug_filename = output_filename.replace('.docx', '.tex')
        debug_tex_path = os.path.join(output_dir, debug_filename)
        with open(debug_tex_path, 'w', encoding='utf-8') as debug_f:
            for fpath in final_files:
                with open(fpath, 'r', encoding='utf-8') as part_f:
                    debug_f.write(part_f.read() + "\n")
        
        try:
            pypandoc.convert_file(
                debug_tex_path,
                'docx',
                format='latex',
                outputfile=output_path,
                extra_args=extra_args
            )
            print(f"  Successfully created {output_path}")
        except Exception as e:
            print(f"  Error: {e}")
            sys.exit(1)
        finally:
            for t in temp_files:
                try: os.remove(t)
                except: pass
        
        return output_path

def post_process_docx(docx_path):
    """
    Applies publisher styles to the generated DOCX file.
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches, Mm
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    except ImportError:
        print("  Warning: python-docx not installed. Skipping post-processing style application.")
        return

    print(f"  Applying Publisher Styles to {docx_path}...")
    doc = Document(docx_path)

    # 1. Page Layout (A4, 0.79" margins)
    # A4 size: 210mm x 297mm
    section = doc.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    
    margin_size = Inches(0.79)
    section.top_margin = margin_size
    section.bottom_margin = margin_size
    section.left_margin = margin_size
    section.right_margin = margin_size

    # 1.5 Fix Numbering Consistency (Force Tab suffix for all levels)
    try:
        if hasattr(doc.part, 'numbering_part') and doc.part.numbering_part:
            numbering_element = doc.part.numbering_part.element
            # Iterate through all abstractNum definitions and their levels
            # We want to enforce <w:suff w:val="tab"/>
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            
            # Find all <w:lvl> elements (levels in abstract numbering)
            for lvl in numbering_element.xpath('.//w:lvl'):
                # Find or create <w:suff>
                suff = lvl.find(qn('w:suff'))
                if suff is None:
                    suff = OxmlElement('w:suff')
                    lvl.append(suff)
                suff.set(qn('w:val'), 'space') # Changed from 'tab' to 'space' per user request
            print("  Enforced Space suffix for numbering levels.")
    except Exception as e:
        print(f"  Warning: Could not patch numbering XML: {e}")

    # 1.6 Add Page Numbers to Footer
    # We need to add a footer to the section and insert a PAGE field.
    # Simple page number centering.
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # We need to insert the page number field using Oxml
    # <w:fldSimple w:instr=" PAGE "/>
    fldSimple = OxmlElement('w:fldSimple')
    fldSimple.set(qn('w:instr'), ' PAGE ')
    # Retrieve the run to put it in, or just append to paragraph
    run = footer_para.add_run()
    run.font.name = "Lora"
    run.font.size = Pt(10)
    # Append the field to the paragraph element, wrapped in a run? 
    # Actually fldSimple contains runs. 
    # Standard way:
    run._element.append(fldSimple)

    def update_style(style_name, font_name="Lora", font_size=11, bold=False, italic=False, 
                     space_before=0, space_after=10, align=None, color_rgb=RGBColor(0,0,0)):
        try:
            style = doc.styles[style_name]
            # Font properties exist on both Character and Paragraph styles (mostly)
            if hasattr(style, 'font'):
                style.font.name = font_name
                # Force rFonts to ensure Theme fallback (like Aptos) doesn't win
                try:
                    rPr = style.element.get_or_add_rPr()
                    # Creating w:rFonts element
                    rFonts = rPr.find(qn('w:rFonts'))
                    if rFonts is None:
                        rFonts = OxmlElement('w:rFonts')
                        rPr.append(rFonts)
                    # Set all types to the target font
                    rFonts.set(qn('w:ascii'), font_name)
                    rFonts.set(qn('w:hAnsi'), font_name)
                    rFonts.set(qn('w:eastAsia'), font_name) # Added
                    rFonts.set(qn('w:cs'), font_name)
                    rFonts.set(qn('w:asciiTheme'), '') # Clear theme bindings
                    rFonts.set(qn('w:hAnsiTheme'), '')
                    rFonts.set(qn('w:eastAsiaTheme'), '')
                    rFonts.set(qn('w:cstheme'), '')
                except Exception as xml_e:
                    print(f"    Warning setting rFonts for {style_name}: {xml_e}")

                style.font.size = Pt(font_size)
                style.font.bold = bold
                style.font.italic = italic
                style.font.color.rgb = color_rgb
            
            # Paragraph formatting only tests on Paragraph styles
            if hasattr(style, 'paragraph_format'):
                style.paragraph_format.space_before = Pt(space_before)
                style.paragraph_format.space_after = Pt(space_after)
                style.paragraph_format.line_spacing = 1.15
                if align is not None:
                    style.paragraph_format.alignment = align
        except KeyError:
            pass

    # Update Styles
    # Normal: Lora, 11pt, Justified
    update_style('Normal', font_name="Lora", font_size=11, space_after=10, align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY)
    
    # Heading 1 (General use in doc, maybe for Chapter Title repeats): Lora 24 Bold
    update_style('Heading 1', font_name="Lora", font_size=24, bold=True, space_before=24, space_after=24)
    
    # Heading 2 (Section): Lora 20 Bold
    # Increased spacing per user request: "start of new section add space before... end of each section add space after"
    # Interpreted as increased margins around the header.
    update_style('Heading 2', font_name="Lora", font_size=20, bold=True, space_before=42, space_after=18)
    
    # Heading 3 (Subsection): Lora 18 Bold (User request: 18pt)
    update_style('Heading 3', font_name="Lora", font_size=18, bold=True, space_before=24, space_after=12, align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY)

    # Heading 4 (Sub-subsection): Lora 16 Bold (User request: 16pt)
    update_style('Heading 4', font_name="Lora", font_size=16, bold=True, space_before=12, space_after=12, align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY)

    # Heading 5: Lora 14 Bold (User request: 14pt)
    update_style('Heading 5', font_name="Lora", font_size=14, bold=True, space_before=12, space_after=12, align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY)

    # Code Blocks (Source Code): Left Aligned
    # Pandoc usually maps verbatims to "Source Code"
    update_style('Source Code', font_name="Consolas", font_size=10, align=WD_PARAGRAPH_ALIGNMENT.LEFT)
    update_style('VerbatimChar', font_name="Consolas", font_size=10, align=WD_PARAGRAPH_ALIGNMENT.LEFT)

    # Captions: Lora 9, Italic, Center
    update_style('Caption', font_name="Lora", font_size=9, italic=True, align=WD_PARAGRAPH_ALIGNMENT.CENTER)
    
    # Also Figure/Table-specific captions often use specific styles or direct formatting.
    # We can try to catch 'Figure Caption' or 'Table Caption' if they exist.
    update_style('Figure Caption', font_name="Lora", font_size=9, italic=True, align=WD_PARAGRAPH_ALIGNMENT.CENTER)
    update_style('Table Caption', font_name="Lora", font_size=9, italic=True, align=WD_PARAGRAPH_ALIGNMENT.CENTER)

    # 2. Custom Title Page Formatting ("CHAPTER X" and Title)
    # R54, G95, B145 Color (User correction: "Red=54, Green=95,Blue=145")
    accent_color = RGBColor(54, 95, 145)

    # We also need to remove the DUPLICATE chapter title that Pandoc generates.
    # The structure is: Custom Title Page (Heading 1 or Normal) -> ... -> \chapter{Title} (Heading 1)
    # Let's clean up the duplicate first.
    # Attempt to find "Heading 1" paragraphs that match the Chapter Title text (assumed known or guessed).
    # Since we don't have the title string here, we look for "Heading 1" paragraphs that appear AFTER the custom title.
    # Our custom title uses explicit formatting (Text 'CHAPTER X' and 'Title'). 
    # The duplicate usually immediately follows the custom title page content.
    
    # Let's perform the formatting of the CUSTOM title page first.
    found_chapter_num = False
    paragraphs_to_remove = []
    
    for p in doc.paragraphs[:20]:
        txt = p.text.strip()
        if not txt: continue
        
        # Check for our custom "CHAPTER X"
        if re.match(r"^CHAPTER \d+$", txt, re.IGNORECASE) and not found_chapter_num:
            # Format Chapter Num
            p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            for run in p.runs:
                run.font.name = "Lora"
                run.font.size = Pt(35)
                run.font.bold = False
                run.font.color.rgb = accent_color
            found_chapter_num = True
            continue
        
        if found_chapter_num:
            # The next non-empty paragraph IS our custom title.
            # Format Title: Lora, 40, Bold, Right
            p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            for run in p.runs:
                run.font.name = "Lora"
                run.font.size = Pt(40)
                run.font.bold = True
                run.font.color.rgb = accent_color
            found_chapter_num = False # Done formatting custom title
            
            # NOW: We expect a Duplicate Title to appear shortly after.
            # It will be a Heading 1 with the same text (roughly) or "1 Title".
            # We will mark the NEXT Heading 1 for removal if it seems redundant.
            # Let's look ahead.
            continue
            
    # Iterate again to removing duplicate Heading 1s
    # Strategy: Remove the FIRST Heading 1 found in the document? 
    # Pandoc `\chapter` creates the first structural Heading 1.
    # Our custom text was likely "Right Aligned" but might have "Normal" style or similar.
    # If we remove the FIRST Heading 1, we remove the structural node.
    # Better: Change the "Structure" Heading 1 to be Hidden? Or just remove text?
    # User asked to "do not need this duplicate line".
    # We will remove the paragraph.
    
    first_h1_removed = False
    for p in doc.paragraphs:
        if p.style.name == 'Heading 1' and not first_h1_removed:
            # Check if this is likely the duplicate
            # It usually contains the Chapter Title.
            # We'll just delete it.
            # deleting a paragraph in python-docx: get parent and remove child
            p._element.getparent().remove(p._element)
            print("  Removed duplicate Heading 1 paragraph.")
            first_h1_removed = True
            break

    # Pre-pass: Identify the LAST "Conclusion" subsection to promote
    target_conclusion_index = -1
    for i, p in enumerate(doc.paragraphs):
        if p.style.name == 'Heading 3':
            # Match "1.2.3 Conclusion" or "Conclusion" with spaces/tabs
            if re.search(r'^\s*[\d\.]+\s+Conclusion\s*$', p.text, re.IGNORECASE) or p.text.strip().lower() == "conclusion":
                target_conclusion_index = i

    # 3. Center Align Images and handle captions if manual
    # Also iterate to apply Aggressive Font/Spacing Fixes for Headings
    for i, paragraph in enumerate(doc.paragraphs):
        # Check for images (drawings)
        if 'w:drawing' in paragraph._element.xml or 'w:pict' in paragraph._element.xml:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
        # Check if style name contains "Caption"
        if "Caption" in paragraph.style.name:
             paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
             for run in paragraph.runs:
                 run.font.name = "Lora"
                 run.font.size = Pt(9)
                 run.font.italic = True
        
        # Special Case: Promote "Conclusion" subsection to Section (Unnumbered)
        # Only for the LAST instance found in the chapter
        if i == target_conclusion_index:
            print(f"  Promoting Conclusion subsection: {paragraph.text.strip()}")
            paragraph.style = doc.styles['Heading 2']
            paragraph.text = "Conclusion"
            try:
                if paragraph._p.pPr is not None and paragraph._p.pPr.numPr is not None:
                     paragraph._p.pPr.remove(paragraph._p.pPr.numPr)
            except Exception:
                pass
        
        # Aggressive Heading Fixes (Spacing and Font)
        if paragraph.style.name in ['Heading 2', 'Heading 3', 'Heading 1', 'Heading 4']:
            # 1. Fix Tab Spacing in Text (Manually numbered by Pandoc?)
            if '\t' in paragraph.text:
                # Replace tab with space
                # We do this by modifying p.text, which consolidates runs. 
                # This is okay because we will re-apply font formatting immediately after.
                paragraph.text = paragraph.text.replace('\t', ' ')
                print(f"  Fixed spacing (Tab->Space) in: {paragraph.text[:30]}...")
            
            # 2. Force Font on Runs (Aptos Override)
            for run in paragraph.runs:
                run.font.name = "Lora"
                # Re-apply bold?
                run.font.bold = True
                
                if paragraph.style.name == 'Heading 1':
                    run.font.size = Pt(24)
                elif paragraph.style.name == 'Heading 2':
                    run.font.size = Pt(20)
                elif paragraph.style.name == 'Heading 3':
                    run.font.size = Pt(18)
                elif paragraph.style.name == 'Heading 4':
                    run.font.size = Pt(16)
                elif paragraph.style.name == 'Heading 5':
                    run.font.size = Pt(14)
                    
                # Explicitly clear Theme fonts in XML if possible, 
                # but run.font.name usually writes w:rFonts w:ascii="Lora" etc.
                if hasattr(run.element.rPr, 'rFonts'):
                     rFonts = run.element.rPr.rFonts
                     if rFonts is not None:
                         rFonts.set(qn('w:asciiTheme'), '')
                         rFonts.set(qn('w:hAnsiTheme'), '')
                         rFonts.set(qn('w:eastAsiaTheme'), '')
                         rFonts.set(qn('w:cstheme'), '')

    # 1.7 Fix Code Block Borders (Optional, if needed)

    # 1.8 Color [FIGURE DETAIL] paragraphs Red (User Request)
    # Iterate through all paragraphs in the document
    try:

        from docx.shared import RGBColor, Pt
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

        red_color = RGBColor(255, 0, 0)
        
        for para in doc.paragraphs:
            txt = para.text.strip()
            if not txt: continue
            
            # Helper to clear and set run text
            def reset_para_style(paragraph, text, font_name="Lora", font_size=11, color=None, bold=False, italic=False):
                # Clear existing runs
                p_element = paragraph._element
                for r in list(paragraph.runs): # List copy to avoid mutation issues?
                    p_element.remove(r._element)
                
                # Add new run
                new_run = paragraph.add_run(text)
                new_run.font.name = font_name
                new_run.font.size = Pt(font_size)
                new_run.font.bold = bold
                new_run.font.italic = italic
                if color:
                    new_run.font.color.rgb = color
                return new_run

            if "[FIGURE DETAIL]" in txt:
                clean_text = txt.replace("[FIGURE DETAIL]", "").strip()
                reset_para_style(para, clean_text, color=red_color, font_size=11)
            
            elif "[GITHUB SCRIPT PROMPT]" in txt:
                clean_text = txt.replace("[GITHUB SCRIPT PROMPT]", "").strip()
                # Remove latex escape if present?
                if clean_text.startswith("\\%"):
                    clean_text = "%" + clean_text[2:]
                
                reset_para_style(para, clean_text, color=red_color, font_size=10, font_name="Consolas")
            
            elif "[LISTING CAPTION]" in txt:
                # Format: Lora 9 Italic throughout, with "Full Listing X.Y" in Bold
                clean_text = txt.replace("[LISTING CAPTION]", "").strip()
                # Ensure we strip any leading # if it somehow made it through
                clean_text = clean_text.lstrip('#').strip()
                
                # We need mixed formatting: Bold part, then Italic part.
                # Regex for "Full Listing X.Y" or "Full Listing 4.1"
                match = re.search(r'^(Full Listing \d+(?:\.\d+)?)(.*)', clean_text)
                
                # Clear runs first
                for r in list(para.runs): para._element.remove(r._element)
                
                para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                if match:
                    bold_part = match.group(1)
                    rest_part = match.group(2)
                    
                    r1 = para.add_run(bold_part)
                    r1.font.name = "Lora"
                    r1.font.size = Pt(9)
                    r1.font.bold = True
                    r1.font.italic = True
                    
                    r2 = para.add_run(rest_part)
                    r2.font.name = "Lora"
                    r2.font.size = Pt(9)
                    r2.font.bold = False
                    r2.font.italic = True
                else:
                    # No formatting match, apply base style
                    r = para.add_run(clean_text)
                    r.font.name = "Lora"
                    r.font.size = Pt(9)
                    r.font.italic = True

            elif "[FIGURE CAPTION]" in txt:
                # Format: "Figure 4.1 - Title..."
                # Entire text Lora 9 Italic
                # "Figure X.Y" part Bold
                
                clean_text = txt.replace("[FIGURE CAPTION]", "").strip()
                
                # Regex for "Figure X.Y" or "Figure 4.1"
                match = re.search(r'^(Figure \d+(?:\.\d+)?)(.*)', clean_text)
                
                # Clear runs
                for r in list(para.runs): para._element.remove(r._element)
                
                para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                if match:
                    bold_part = match.group(1)
                    rest_part = match.group(2)
                    
                    r1 = para.add_run(bold_part)
                    r1.font.name = "Lora"
                    r1.font.size = Pt(9)
                    r1.font.bold = True
                    r1.font.italic = True
                    
                    r2 = para.add_run(rest_part)
                    r2.font.name = "Lora"
                    r2.font.size = Pt(9)
                    r2.font.bold = False
                    r2.font.italic = True
                else:
                    r = para.add_run(clean_text)
                    r.font.name = "Lora"
                    r.font.size = Pt(9)
                    r.font.italic = True


        print("  Applied Publisher Styles and Formatted Details.")

    except Exception as e:
        print(f"  Warning: Could not apply publisher styles/details: {e}")

    try:
        doc.save(docx_path)
        print("  Publisher Styles applied successfully.")
    except Exception as e:
        print(f"  Error saving styled DOCX: {e}")


def main():
    parser = argparse.ArgumentParser(description="Convert LaTeX to Publisher Style Docx.")
    parser.add_argument("--chapter", type=int, help="Specific chapter number to convert.")
    args = parser.parse_args()

    metadata_file = "input/metadata.json"
    output_dir = "output"
    
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        
    output_path = convert_book(metadata_file, output_dir, args.chapter)
    
    if output_path:
        post_process_docx(output_path)

if __name__ == "__main__":
    main()
